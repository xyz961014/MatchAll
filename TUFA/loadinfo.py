#coding=utf-8
from docx import Document
import pymysql.cursors
import re

def loadteamandleaderinfomanan(filename):
  if not re.search('\.docx$',filename):
    return
  document = Document(filename)
  ps = document.paragraphs
  teaminfo = ps[2].text.split()
  print(teaminfo)
  tnObj = re.match(r'院系：\S',teaminfo[0])
  col = 0
  if tnObj:
    teamname = teaminfo[0][3:]
    col += 1
  else:
    teamname = teaminfo[1]
    col += 2
  colObj = re.match(r'比赛服颜色（上衣/短裤）：\S',teaminfo[col])
  if colObj:
    teamcolor = teaminfo[col][13:]
  else:
    teamcolor = teaminfo[col+1]
  print(teamname,teamcolor)
  tb = document.tables
  connection = pymysql.connect(host='localhost',
                               user='root',
                               password='961014',
                               db='MANAN_1718',
                               charset='utf8mb4',
                               cursorclass=pymysql.cursors.DictCursor)

  try:
    if len(tb) > 1:
      for p in tb[1].columns:
        if p.cells[0].text != '姓名':
          name = p.cells[0].text
          job = p.cells[1].text
          phonenumber = p.cells[2].text
          email = p.cells[3].text
          with connection.cursor() as cursor:
          # Create a new record
            sql = "INSERT INTO `Leaders` (`Team`, `Name`,`Job`,`PhoneNumber`,`Email`) VALUES (%s, %s, %s, %s, %s)"
            cursor.execute(sql, (teamname, name, job, phonenumber, email))
    with connection.cursor() as cursor:
      sql = "INSERT INTO `Teams` (`TeamName`, `KitColor`) VALUES (%s, %s)"
      cursor.execute(sql,(teamname,teamcolor))

    # connection is not autocommit by default. So you must commit to save
    # your changes.
      connection.commit()

  finally:
    connection.close()


def loadplayerinfomanan(filename):
  if not re.search('\.docx$',filename):
    return
  document = Document(filename)
  ps = document.paragraphs
  print(len(ps))
  teaminfo = ps[2].text.split()
  print(teaminfo)
  tnObj = re.match(r'院系：\S',teaminfo[0])
  col = 0
  if tnObj:
    teamname = teaminfo[0][3:]
    col += 1
  else:
    teamname = teaminfo[1]
    col += 2
  colObj = re.match(r'比赛服颜色（上衣/短裤）：\S',teaminfo[col])
  if colObj:
    teamcolor = teaminfo[col][13:]
  else:
    teamcolor = teaminfo[col+1]
  print(teamname,teamcolor)
  tb = document.tables
  connection = pymysql.connect(host='localhost',
                               user='root',
                               password='961014',
                               db='MANAN_1718',
                               charset='utf8mb4',
                               cursorclass=pymysql.cursors.DictCursor)

  try:
    for p in tb[0].rows[1:]:
      name = p.cells[1].text
      class1 = p.cells[2].text
      idnumber = p.cells[3].text
      phonenumber = p.cells[4].text
      kitnumber = p.cells[5].text
      extrainfo = p.cells[6].text
      if not (re.match('^\s+$', name) or name == ''):
        name = re.sub(r'^\s+', '', name)
        name = re.sub(r'\s+$', '', name)
        kitnumber = int(kitnumber)
        if re.match('^\s+$',idnumber) or idnumber == '无' or idnumber == '':
          idnumber = None
        with connection.cursor() as cursor:
          # Create a new record
          sql = "INSERT INTO `Players` (`Team`, `Name`,`Class`,`IDNumber`,`PhoneNumber`,`KitNumber`,`ExtraInfo`) VALUES (%s, %s, %s, %s, %s, %s, %s)"
          cursor.execute(sql, (teamname, name, class1, idnumber, phonenumber, kitnumber, extrainfo))

    # connection is not autocommit by default. So you must commit to save
    # your changes.
      connection.commit()

  finally:
    connection.close()
  
def loadinfofreshman(filename):
    if not re.search('\.docx$',filename):
        return
    document = Document(filename)
    ps = document.paragraphs
    #print(len(ps))
    teaminfo = ps[1].text.split()
    #print(teaminfo)
    tnObj = re.match(r'球队名称：\S',teaminfo[0])
    col = 0
    if tnObj:
      teamname = teaminfo[0][5:]
      col += 1
    else:
      teamname = teaminfo[1]
      col += 2
    colObj = re.match(r'比赛服颜色（上衣/短裤）：\S',teaminfo[col])
    if colObj:
      teamcolor = teaminfo[col][13:]
    else:
      teamcolor = teaminfo[col+1]
    print(teamname,teamcolor)
    tb = document.tables
    connection = pymysql.connect(host='localhost',
                                 user='root',
                                 password='961014',
                                 db='FRESHMANCUP_18',
                                 charset='utf8mb4',
                                 cursorclass=pymysql.cursors.DictCursor)

    try:
      for p in tb[1].rows[1:]:
        name = p.cells[1].text
        class1 = p.cells[3].text
        idnumber = p.cells[4].text
        phonenumber = p.cells[5].text
        kitnumber = p.cells[6].text
        extrainfo = p.cells[7].text
        #print(name, class1,idnumber,phonenumber,kitnumber,extrainfo)
        if not (re.match('^\s+$', name) or name == ''):
          name = re.sub(r'^\s+', '', name)
          name = re.sub(r'\s+$', '', name)
          kitnumber = int(kitnumber)
          if re.match('^\s+$',idnumber) or idnumber == '无' or idnumber == '':
            idnumber = None
          print(name, class1, idnumber, phonenumber, kitnumber, extrainfo)
          with connection.cursor() as cursor:
            # Create a new record
            sql = "INSERT INTO `Players` (`Team`, `Name`,`Class`,`IDNumber`,`PhoneNumber`,`KitNumber`,`ExtraInfo`) VALUES (%s, %s, %s, %s, %s, %s, %s)"
            cursor.execute(sql, (teamname, name, class1, idnumber, phonenumber, kitnumber, extrainfo))

      # connection is not autocommit by default. So you must commit to save
      # your changes.
      with connection.cursor() as cursor:
        sql = "INSERT INTO `Teams` (`TeamName`, `KitColor`) VALUES (%s, %s)"
        cursor.execute(sql,(teamname,teamcolor))
      connection.commit()
    finally:
      connection.close()


def loadinfobubaoming(filename):
    if not re.search('\.docx$',filename):
        return
    document = Document(filename)
    ps = document.paragraphs
    teaminfo = ps[3].text.split()
    tnObj = re.match(r'参赛队：\S',teaminfo[0])
    col = 0
    if tnObj:
      teamname = teaminfo[0][4:]
    else:
      teamname = teaminfo[1]
    print(teamname)
    tb = document.tables
    connection = pymysql.connect(host='localhost',
                                user='root',
                                password='961014',
                                db='MANAN_1718',
                                charset='utf8mb4',
                                cursorclass=pymysql.cursors.DictCursor)
    try:
        with connection.cursor() as cursor:
            sql = 'SELECT * FROM Players WHERE Team = %s'
            cursor.execute(sql,teamname)
            tp = cursor.fetchall()
            nums = [i['KitNumber'] for i in tp]
            numtp = len(tp)
            for p in tb[0].rows[1:]:
                if not (re.match('\s+', p.cells[0].text) or p.cells[0].text == ''):
                    numtp -= 1
                    print(numtp)
            for p in tb[1].rows[1:]:
                if not (re.match('\s+', p.cells[1].text) or p.cells[1].text == ''):
                    numtp += 1
                    print(numtp)
            if numtp > 30:
                raise NameError('报名人数不符合规定')
        for p in tb[0].rows[1:]:
            name = p.cells[1].text
            class1 = p.cells[2].text
            idnumber = p.cells[3].text
            phonenumber = p.cells[4].text
            kitnumber = p.cells[5].text
            extrainfo = p.cells[6].text
            if not (re.match('^\s+$', name) or name == ''):
                name = re.sub(r'^\s+', '', name)
                name = re.sub(r'\s+$', '', name)
                kitnumber = int(kitnumber)
                nums.remove(kitnumber)
                if re.match('^\s+$',idnumber) or idnumber == '无' or idnumber == '':
                    idnumber = None
                print(name,class1,idnumber,phonenumber,kitnumber,extrainfo)
                with connection.cursor() as cursor:
                    sql = "UPDATE `Players` SET Valid = 0 WHERE Team = %s AND Name = %s AND IDNumber = %s AND KitNumber = %s"
                    cursor.execute(sql, (teamname, name, idnumber, kitnumber))


        for p in tb[1].rows[1:]:
            name = p.cells[1].text
            class1 = p.cells[2].text
            idnumber = p.cells[3].text
            phonenumber = p.cells[4].text
            kitnumber = p.cells[5].text
            extrainfo = p.cells[6].text
            if not (re.match('^\s+$', name) or name == ''):
                name = re.sub(r'^\s+', '', name)
                name = re.sub(r'\s+$', '', name)
                if int(kitnumber) in nums:
                    raise NameError("有重复的球员号码")
                else:
                    nums.append(int(kitnumber))
                kitnumber = int(kitnumber)
                if re.match('^\s+$',idnumber) or idnumber == '无' or idnumber == '':
                    idnumber = None
                print(name,class1,idnumber,phonenumber,kitnumber,extrainfo)
                with connection.cursor() as cursor:
                    sql = "INSERT INTO `Players` (`Team`, `Name`,`Class`,`IDNumber`,`PhoneNumber`,`KitNumber`,`ExtraInfo`) VALUES (%s, %s, %s, %s, %s, %s, %s)"
                    cursor.execute(sql, (teamname, name, class1, idnumber, phonenumber, kitnumber, extrainfo))
        connection.commit()
    finally:
        connection.close()

def loadinfomanyu(filename):
    if not re.search('\.docx$',filename):
        return
    document = Document(filename)
    ps = document.paragraphs
    print(len(ps))
    teaminfo = ps[2].text.split()
    print(teaminfo)
    tnObj = re.match(r'院系：\S',teaminfo[0])
    col = 0
    if tnObj:
        teamname = teaminfo[0][3:]
        col += 1
    else:
        teamname = teaminfo[1]
        col += 2
    colObj = re.match(r'比赛服颜色（上衣/短裤）：\S',teaminfo[col])
    if colObj:
        teamcolor = teaminfo[col][13:]
    else:
        teamcolor = teaminfo[col+1]
    print(teamname,teamcolor)
    tb = document.tables
    connection = pymysql.connect(host='localhost',
                                 user='root',
                                 password='961014',
                                 db='MANYU_18',
                                 charset='utf8mb4',
                                 cursorclass=pymysql.cursors.DictCursor)

    try:
        for p in tb[0].rows[1:]:
            name = p.cells[1].text
            phonenumber = p.cells[2].text
            job = p.cells[3].text
        print(name,phonenumber,job)
        with connection.cursor() as cursor:
            sql = "INSERT INTO `Leaders` (`Team`, `Name`,`Job`,`PhoneNumber`) VALUES (%s, %s, %s, %s)"
            cursor.execute(sql, (teamname, name, job, phonenumber))
        with connection.cursor() as cursor:
            sql = "INSERT INTO `Teams` (`TeamName`, `KitColor`) VALUES (%s, %s)"
            cursor.execute(sql,(teamname,teamcolor))
        kitnums = []
        for p in tb[1].rows[1:]:
            name = p.cells[2].text
            class1 = p.cells[3].text
            idnumber = p.cells[4].text
            phonenumber = p.cells[5].text
            kitnumber = p.cells[1].text
            extrainfo = p.cells[6].text
            if not (re.match('^\s+$', name) or name == ''):
                name = re.sub(r'^\s+', '', name)
                name = re.sub(r'\s+$', '', name)
                if int(kitnumber) in kitnums:
                    raise NameError('有重复的球员号码')
                else:
                    kitnums.append(int(kitnumber))
                kitnumber = int(kitnumber)
                print(name,class1,idnumber,phonenumber,kitnumber,extrainfo)
                if re.match('^\s+$',idnumber) or idnumber == '无' or idnumber == '':
                    idnumber = None
                with connection.cursor() as cursor:
                    # Create a new record
                    sql = "INSERT INTO `Players` (`Team`, `Name`,`Class`,`IDNumber`,`PhoneNumber`,`KitNumber`,`ExtraInfo`) VALUES (%s, %s, %s, %s, %s, %s, %s)"
                    cursor.execute(sql, (teamname, name, class1, idnumber, phonenumber, kitnumber, extrainfo))

        # connection is not autocommit by default. So you must commit to save
        # your changes.
        connection.commit()

    finally:
        connection.close()

def loadinfomawu(filename):
    if not re.search('\.docx$',filename):
        return
    document = Document(filename)
    ps = document.paragraphs
    print(len(ps))
    teaminfo = ps[3].text.split()
    teamname = teaminfo[1]
    teamcolor = teaminfo[3] + '/' + teaminfo[5]
    print(teamname, teamcolor)
    tb = document.tables
    connection = pymysql.connect(host='localhost',
                                 user='root',
                                 password='961014',
                                 db='MAWU_18',
                                 charset='utf8mb4',
                                 cursorclass=pymysql.cursors.DictCursor)

    try:
        with connection.cursor() as cursor:
            sql = "INSERT INTO `Teams` (`TeamName`, `KitColor`) VALUES (%s, %s)"
            cursor.execute(sql,(teamname,teamcolor))
        kitnums = []
        for p in tb[0].rows[1:]:
            name = p.cells[2].text
            class1 = p.cells[3].text
            idnumber = p.cells[4].text
            phonenumber = p.cells[5].text
            kitnumber = p.cells[1].text
            extrainfo = p.cells[6].text
            if not (re.match('^\s+$', name) or name == ''):
                name = re.sub(r'^\s+', '', name)
                name = re.sub(r'\s+$', '', name)
                if int(kitnumber) in kitnums:
                    raise NameError('有重复的球员号码')
                else:
                    kitnums.append(int(kitnumber))
                kitnumber = int(kitnumber)
                print(name,class1,idnumber,phonenumber,kitnumber,extrainfo)
                if re.match('\s+',idnumber) or idnumber == '无' or idnumber == '':
                    idnumber = None
                with connection.cursor() as cursor:
                    # Create a new record
                    sql = "INSERT INTO `Players` (`Team`, `Name`,`Class`,`IDNumber`,`PhoneNumber`,`KitNumber`,`ExtraInfo`) VALUES (%s, %s, %s, %s, %s, %s, %s)"
                    cursor.execute(sql, (teamname, name, class1, idnumber, phonenumber, kitnumber, extrainfo))

        # connection is not autocommit by default. So you must commit to save
        # your changes.
        #connection.commit()

    finally:
        connection.close()



#loadinfofreshman('./新生杯队伍报名表/新生杯清北人叉队报名表.docx')
#loadinfobubaoming('./男足补报名18/电机系马杯男足补报名表.docx')
#loadinfomanyu('./女足报名18/template.docx')
#loadinfomawu('./马五2018/化学系马杯五人制报名材料.docx')
#loadinfomawu('./马五2018/土木院（系）马杯五人制报名材料.docx')
