#coding=utf-8
from docx import Document
from docx.shared import Pt
import pymysql.cursors
import re
import TeamDict

def createmanansheet(hometeam='',awayteam='',title='2018—2019年度清华大学马约翰学生运动会',matchdb = 'MANAN_1718',subtitle = '男子足球执场单'):
    filename = '/var/www/TUFA/sheettemplate.docx'
    doc = Document(filename)
    paras = doc.paragraphs
    tables = doc.tables
    styles = doc.styles
    #renew the title
    paras[0].runs[0].text = title
    for rest in paras[0].runs[1:]:
        rest.text = ''
    paras[13].runs[0].text = title
    for rest in paras[13].runs[1:]:
        rest.text = ''
    paras[26].runs[0].text = title
    for rest in paras[26].runs[1:]:
        rest.text = ''

    paras[1].runs[0].text =subtitle
    for rest in paras[1].runs[1:]:
        rest.text = ''
    paras[14].runs[0].text = subtitle
    for rest in paras[14].runs[1:]:
        rest.text = ''
    paras[27].runs[0].text = subtitle
    for rest in paras[27].runs[1:]:
        rest.text = ''
    
    homeplayers = []
    awayplayers = []
    homeleaders = []
    awayleaders = []
    homeinfo = []
    awayinfo = []
    connection = pymysql.connect(host='localhost',
                                 user='root',
                                 password='961014',
                                 db=matchdb,
                                 charset='utf8mb4',
                                 cursorclass=pymysql.cursors.DictCursor)
    

    
    try:
        with connection.cursor() as cursor:
            sql = "SELECT * FROM Players WHERE Team = %s AND Valid = 1 ORDER BY KitNumber"
            cursor.execute(sql,hometeam)
            homeplayers = cursor.fetchall()
            cursor.execute(sql,awayteam)
            awayplayers = cursor.fetchall()
            sql = "SELECT * FROM Leaders WHERE Team = %s"
            cursor.execute(sql,hometeam)
            homeleaders = cursor.fetchall()
            cursor.execute(sql,awayteam)
            awayleaders = cursor.fetchall()
            sql = "SELECT * FROM Teams WHERE TeamName = %s"
            cursor.execute(sql,hometeam)
            homeinfo = cursor.fetchall()
            cursor.execute(sql,awayteam)
            awayinfo = cursor.fetchall()

    finally:
        connection.close()
    #print(len(tables))
    #for index,p in enumerate(tables):
     # print('index:',index)
      #for iq,q in enumerate(p.rows):
       # print('row:',iq)
        #for iw,w in enumerate(q.cells):
         # print('cell',iw,":",w.text) 
    tables[0].rows[0].cells[0].paragraphs[0].runs[0].text = homeinfo[0]['TeamName']
    tables[0].rows[0].cells[4].paragraphs[0].runs[1].text += homeinfo[0]['KitColor']
    tables[16].rows[0].cells[0].paragraphs[0].runs[0].text = homeinfo[0]['TeamName']
    tables[16].rows[0].cells[3].paragraphs[0].runs[1].text += homeinfo[0]['KitColor']
    if len(homeleaders) > 0:
        tables[0].rows[32].cells[2].paragraphs[0].runs[0].text = homeleaders[0]['Name']
        tables[0].rows[32].cells[5].paragraphs[0].runs[0].text = homeleaders[1]['Name']
    lenhc = len(tables[0].rows[0].cells[4].paragraphs[0].runs[0].text) + len(tables[0].rows[0].cells[4].paragraphs[0].runs[1].text)
    if lenhc >= 10:
        tables[0].rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(110/lenhc)
        tables[0].rows[0].cells[4].paragraphs[0].runs[1].font.size = Pt(110/lenhc)
        tables[16].rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(110/lenhc)
        tables[16].rows[0].cells[3].paragraphs[0].runs[1].font.size = Pt(110/lenhc)
      
      
    
    for i in range(len(homeplayers)):
        tables[0].rows[i+2].cells[0].paragraphs[0].runs[0].text = str(homeplayers[i]['KitNumber'])
        tables[16].rows[i+2].cells[0].paragraphs[0].runs[0].text = str(homeplayers[i]['KitNumber'])
        extra = ""
        if re.search('校友', homeplayers[i]['ExtraInfo']):
            extra += "校友"
        if re.search('教工', homeplayers[i]['ExtraInfo']):
            extra += "教工"
        if re.search('足特', homeplayers[i]['ExtraInfo']):
            extra += "足特"
        if re.search('队长', homeplayers[i]['ExtraInfo']):
            extra += "C"
        if not extra == "":
            tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].text = homeplayers[i]['Name'] + '(' + extra + ')'
            tables[16].rows[i+2].cells[1].paragraphs[0].runs[0].text = homeplayers[i]['Name'] + '(' + extra + ')'
        else:
            tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].text = homeplayers[i]['Name']
            tables[16].rows[i+2].cells[1].paragraphs[0].runs[0].text = homeplayers[i]['Name']
        lenname = len(tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].text)
        utf8_lenname = len(tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].text.encode('utf-8'))
        lenname = (utf8_lenname - lenname) / 2 * 1.5 + lenname
        if lenname >= 16:
            tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].font.size = Pt(146/lenname)
            tables[16].rows[i+2].cells[1].paragraphs[0].runs[0].font.size = Pt(146/lenname)
        if homeplayers[i]['Suspension']:
            tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].font.strike = True
            tables[16].rows[i+2].cells[1].paragraphs[0].runs[0].font.strike = True
        
     
    tables[8].rows[0].cells[0].paragraphs[0].runs[0].text = awayinfo[0]['TeamName']
    tables[8].rows[0].cells[4].paragraphs[0].runs[1].text += awayinfo[0]['KitColor']
    tables[16].rows[0].cells[6].paragraphs[0].runs[0].text = awayinfo[0]['TeamName']
    tables[16].rows[0].cells[9].paragraphs[0].runs[1].text += awayinfo[0]['KitColor']
    if len(awayleaders) > 0:
        tables[8].rows[32].cells[2].paragraphs[0].runs[0].text = awayleaders[0]['Name']
        tables[8].rows[32].cells[5].paragraphs[0].runs[0].text = awayleaders[1]['Name']
    lenac = len(tables[8].rows[0].cells[4].paragraphs[0].runs[0].text) + len(tables[8].rows[0].cells[4].paragraphs[0].runs[1].text)
    if lenac >= 10:
        tables[8].rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(110/lenac)
        tables[8].rows[0].cells[4].paragraphs[0].runs[1].font.size = Pt(110/lenac)
        tables[16].rows[0].cells[9].paragraphs[0].runs[0].font.size = Pt(110/lenac)
        tables[16].rows[0].cells[9].paragraphs[0].runs[1].font.size = Pt(110/lenac)

    for i in range(len(awayplayers)):
        tables[8].rows[i+2].cells[0].paragraphs[0].runs[0].text = str(awayplayers[i]['KitNumber'])
        tables[16].rows[i+2].cells[6].paragraphs[0].runs[0].text = str(awayplayers[i]['KitNumber'])
        extra = ""
        if re.search('校友', awayplayers[i]['ExtraInfo']):
            extra += "校友"
        if re.search('教工', awayplayers[i]['ExtraInfo']):
            extra += "教工"
        if re.search('足特', awayplayers[i]['ExtraInfo']):
            extra += "足特"
        if re.search('队长', awayplayers[i]['ExtraInfo']):
            extra += "C"
        if not extra == "":
            tables[8].rows[i+2].cells[1].paragraphs[0].runs[0].text = awayplayers[i]['Name'] + '(' + extra + ')'
            tables[16].rows[i+2].cells[7].paragraphs[0].runs[0].text = awayplayers[i]['Name'] + '(' + extra + ')'
        else:
            tables[8].rows[i+2].cells[1].paragraphs[0].runs[0].text = awayplayers[i]['Name']
            tables[16].rows[i+2].cells[7].paragraphs[0].runs[0].text = awayplayers[i]['Name']
        lenname = len(tables[8].rows[i+2].cells[1].paragraphs[0].runs[0].text)
        utf8_lenname = len(tables[8].rows[i+2].cells[1].paragraphs[0].runs[0].text.encode('utf-8'))
        lenname = (utf8_lenname - lenname) / 2 * 1.5 + lenname
        if lenname >= 16:
            tables[8].rows[i+2].cells[1].paragraphs[0].runs[0].font.size = Pt(146/lenname)
            tables[16].rows[i+2].cells[7].paragraphs[0].runs[0].font.size = Pt(146/lenname)
        if awayplayers[i]['Suspension']:
            tables[8].rows[i+2].cells[1].paragraphs[0].runs[0].font.strike = True
            tables[16].rows[i+2].cells[7].paragraphs[0].runs[0].font.strike = True
      
    fname = '/var/www/TUFA/sheets/' + TeamDict.getabbr(homeinfo[0]['TeamName']) + TeamDict.getabbr(awayinfo[0]['TeamName']) + 'sheet.docx'
    doc.save(fname)
    return fname,fname[21:]

def createsinglesheet(hometeam='',awayteam='',title='2017—2018年度清华大学马约翰学生运动会',matchdb = 'MAWU_18',subtitle = '五人制足球执场单'):
    if re.match('MAWU', matchdb):
        filename = '/var/www/TUFA/templatemawu.docx'
    elif re.match('MANYU', matchdb):
        filename = '/var/www/TUFA/templatemanyu.docx'
    doc = Document(filename)
    paras = doc.paragraphs
    tables = doc.tables
    styles = doc.styles
    #renew the title
    paras[0].runs[0].text = title
    for rest in paras[0].runs[1:]:
      rest.text = ''
    paras[1].runs[0].text =subtitle
    for rest in paras[1].runs[1:]:
      rest.text = ''
    
    homeplayers = []
    awayplayers = []
    homeleaders = []
    awayleaders = []
    homeinfo = []
    awayinfo = []
    connection = pymysql.connect(host='localhost',
                                 user='root',
                                 password='961014',
                                 db=matchdb,
                                 charset='utf8mb4',
                                 cursorclass=pymysql.cursors.DictCursor)
    try:
      with connection.cursor() as cursor:
        sql = "SELECT * FROM Players WHERE Team = %s AND Valid = 1 ORDER BY KitNumber"
        cursor.execute(sql,hometeam)
        homeplayers = cursor.fetchall()
        cursor.execute(sql,awayteam)
        awayplayers = cursor.fetchall()
        sql = "SELECT * FROM Leaders WHERE Team = %s"
        cursor.execute(sql,hometeam)
        homeleaders = cursor.fetchall()
        cursor.execute(sql,awayteam)
        awayleaders = cursor.fetchall()
        sql = "SELECT * FROM Teams WHERE TeamName = %s"
        cursor.execute(sql,hometeam)
        homeinfo = cursor.fetchall()
        cursor.execute(sql,awayteam)
        awayinfo = cursor.fetchall()

    finally:
      connection.close()
    #print(len(tables))
    #for index,p in enumerate(tables):
    #  print('index:',index)
    #  for iq,q in enumerate(p.rows):
    #    print('row:',iq)
    #    for iw,w in enumerate(q.cells):
    #      print('cell',iw,":",w.text) 
    tables[0].rows[0].cells[0].paragraphs[0].runs[0].text = homeinfo[0]['TeamName']
    tables[0].rows[0].cells[3].paragraphs[0].runs[1].text += homeinfo[0]['KitColor']
    #if len(homeleaders) > 0:
    #  tables[0].rows[32].cells[2].paragraphs[0].runs[0].text = homeleaders[0]['Name']
    #  tables[0].rows[32].cells[5].paragraphs[0].runs[0].text = homeleaders[1]['Name']
    lenhc = len(tables[0].rows[0].cells[3].paragraphs[0].runs[0].text) + len(tables[0].rows[0].cells[3].paragraphs[0].runs[1].text)
    if lenhc >= 10:
      tables[0].rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(110/lenhc)
      tables[0].rows[0].cells[3].paragraphs[0].runs[1].font.size = Pt(110/lenhc)
      
      
    
    for i in range(len(homeplayers)):
      tables[0].rows[i+2].cells[0].paragraphs[0].runs[0].text = str(homeplayers[i]['KitNumber'])
      if homeplayers[i]['ExtraInfo'] and not re.match('^\s+$',homeplayers[i]['ExtraInfo']):
        tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].text = homeplayers[i]['Name'] + '(' + homeplayers[i]['ExtraInfo'] + ')'
      else:
        tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].text = homeplayers[i]['Name']
      lenname = len(tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].text)
      utf8_lenname = len(tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].text.encode('utf-8'))
      lenname = (utf8_lenname - lenname) / 2 * 1.5 + lenname
      if lenname >= 16:
        tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].font.size = Pt(146/lenname)
      if homeplayers[i]['Suspension']:
        tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].font.strike = True
        
     
    tables[0].rows[0].cells[6].paragraphs[0].runs[0].text = awayinfo[0]['TeamName']
    tables[0].rows[0].cells[9].paragraphs[0].runs[1].text += awayinfo[0]['KitColor']
    #if len(awayleaders) > 0:
    #  tables[8].rows[32].cells[2].paragraphs[0].runs[0].text = awayleaders[0]['Name']
    #  tables[8].rows[32].cells[5].paragraphs[0].runs[0].text = awayleaders[1]['Name']
    lenac = len(tables[0].rows[0].cells[9].paragraphs[0].runs[0].text) + len(tables[0].rows[0].cells[9].paragraphs[0].runs[1].text)
    if lenac >= 10:
      tables[0].rows[0].cells[9].paragraphs[0].runs[0].font.size = Pt(110/lenac)
      tables[0].rows[0].cells[9].paragraphs[0].runs[1].font.size = Pt(110/lenac)

    for i in range(len(awayplayers)):
      tables[0].rows[i+2].cells[6].paragraphs[0].runs[0].text = str(awayplayers[i]['KitNumber'])
      if awayplayers[i]['ExtraInfo'] and not re.match('\s+',awayplayers[i]['ExtraInfo']):
        tables[0].rows[i+2].cells[7].paragraphs[0].runs[0].text = awayplayers[i]['Name'] + '(' + awayplayers[i]['ExtraInfo'] + ')'
      else:
        tables[0].rows[i+2].cells[7].paragraphs[0].runs[0].text = awayplayers[i]['Name']
      lenname = len(tables[0].rows[i+2].cells[7].paragraphs[0].runs[0].text)
      utf8_lenname = len(tables[0].rows[i+2].cells[7].paragraphs[0].runs[0].text.encode('utf-8'))
      lenname = (utf8_lenname - lenname) / 2 * 1.5 + lenname
      if lenname >= 16:
        tables[0].rows[i+2].cells[7].paragraphs[0].runs[0].font.size = Pt(146/lenname)
      if awayplayers[i]['Suspension']:
        tables[0].rows[i+2].cells[7].paragraphs[0].runs[0].font.strike = True
      
    fname = '/var/www/TUFA/sheets/' + TeamDict.getabbr(homeinfo[0]['TeamName']) + TeamDict.getabbr(awayinfo[0]['TeamName']) + 'sheet.docx'
    doc.save(fname)
    return fname,fname[21:]
    
def create4subsheet(hometeam='',awayteam='',title='2017—2018年度清华大学',matchdb = 'FRESHMANCUP_17',subtitle = '新生杯比赛执场单'):
    filename = '/var/www/TUFA/template4sub.docx'
    doc = Document(filename)
    paras = doc.paragraphs
    tables = doc.tables
    styles = doc.styles
    #renew the title
    paras[0].runs[0].text = title
    for rest in paras[0].runs[1:]:
      rest.text = ''
    paras[1].runs[0].text =subtitle
    for rest in paras[1].runs[1:]:
      rest.text = ''
    
    
    homeplayers = []
    awayplayers = []
    homeleaders = []
    awayleaders = []
    homeinfo = []
    awayinfo = []
    connection = pymysql.connect(host='localhost',
                                 user='root',
                                 password='961014',
                                 db=matchdb,
                                 charset='utf8mb4',
                                 cursorclass=pymysql.cursors.DictCursor)
    try:
      with connection.cursor() as cursor:
        sql = "SELECT * FROM Players WHERE Team = %s AND Valid = 1 ORDER BY KitNumber"
        cursor.execute(sql,hometeam)
        homeplayers = cursor.fetchall()
        cursor.execute(sql,awayteam)
        awayplayers = cursor.fetchall()
        sql = "SELECT * FROM Teams WHERE TeamName = %s"
        cursor.execute(sql,hometeam)
        homeinfo = cursor.fetchall()
        cursor.execute(sql,awayteam)
        awayinfo = cursor.fetchall()

    finally:
      connection.close()
    #print(len(tables))
    #for index,p in enumerate(tables):
    #  print('index:',index)
    #  for iq,q in enumerate(p.rows):
    #    print('row:',iq)
    #    for iw,w in enumerate(q.cells):
    #      print('cell',iw,":",w.text) 
    tables[0].rows[0].cells[0].paragraphs[0].runs[0].text = homeinfo[0]['TeamName']
    tables[0].rows[0].cells[7].paragraphs[0].runs[1].text += homeinfo[0]['KitColor']
    lenhc = len(tables[0].rows[0].cells[7].paragraphs[0].runs[0].text) + len(tables[0].rows[0].cells[7].paragraphs[0].runs[1].text)
    if lenhc >= 10:
      tables[0].rows[0].cells[7].paragraphs[0].runs[0].font.size = Pt(110/lenhc)
      tables[0].rows[0].cells[7].paragraphs[0].runs[1].font.size = Pt(110/lenhc)
      
      
    
    for i in range(len(homeplayers)):
      if not re.match('NANQI', matchdb):
        tables[0].rows[i+2].cells[0].paragraphs[0].runs[0].text = str(homeplayers[i]['KitNumber'])
      if homeplayers[i]['ExtraInfo'] and not re.match('^\s+$',homeplayers[i]['ExtraInfo']):
        tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].text = homeplayers[i]['Name'] + '(' + homeplayers[i]['ExtraInfo'] + ')'
      else:
        tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].text = homeplayers[i]['Name']
      lenname = len(tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].text)
      utf8_lenname = len(tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].text.encode('utf-8'))
      lenname = (utf8_lenname - lenname) / 2 * 1.5 + lenname
      if lenname >= 16:
        tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].font.size = Pt(146/lenname)
      if homeplayers[i]['Suspension']:
        tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].font.strike = True
        
     
    tables[0].rows[0].cells[10].paragraphs[0].runs[0].text = awayinfo[0]['TeamName']
    tables[0].rows[0].cells[16].paragraphs[0].runs[1].text += awayinfo[0]['KitColor']
    lenac = len(tables[0].rows[0].cells[16].paragraphs[0].runs[0].text) + len(tables[0].rows[0].cells[16].paragraphs[0].runs[1].text)
    if lenac >= 10:
      tables[0].rows[0].cells[16].paragraphs[0].runs[0].font.size = Pt(110/lenac)
      tables[0].rows[0].cells[16].paragraphs[0].runs[1].font.size = Pt(110/lenac)

    for i in range(len(awayplayers)):
      if not re.match('NANQI', matchdb):
        tables[0].rows[i+2].cells[10].paragraphs[0].runs[0].text = str(awayplayers[i]['KitNumber'])
      if awayplayers[i]['ExtraInfo'] and not re.match('\s+',awayplayers[i]['ExtraInfo']):
        tables[0].rows[i+2].cells[11].paragraphs[0].runs[0].text = awayplayers[i]['Name'] + '(' + awayplayers[i]['ExtraInfo'] + ')'
      else:
        tables[0].rows[i+2].cells[11].paragraphs[0].runs[0].text = awayplayers[i]['Name']
      lenname = len(tables[0].rows[i+2].cells[11].paragraphs[0].runs[0].text)
      utf8_lenname = len(tables[0].rows[i+2].cells[11].paragraphs[0].runs[0].text.encode('utf-8'))
      lenname = (utf8_lenname - lenname) / 2 * 1.5 + lenname
      if lenname >= 16:
        tables[0].rows[i+2].cells[11].paragraphs[0].runs[0].font.size = Pt(146/lenname)
      if awayplayers[i]['Suspension']:
        tables[0].rows[i+2].cells[11].paragraphs[0].runs[0].font.strike = True
      
    fname = '/var/www/TUFA/sheets/' + 'sheet.docx'
    doc.save(fname)
    return fname,fname[21:]

def create5subsheet(hometeam='',awayteam='',title='2018—2019年度清华大学',matchdb = 'NYUQI_1819',subtitle = '女子七人制联赛比赛执场单'):
    filename = '/var/www/TUFA/template5sub.docx'
    doc = Document(filename)
    paras = doc.paragraphs
    tables = doc.tables
    styles = doc.styles
    #renew the title
    paras[0].runs[0].text = title
    for rest in paras[0].runs[1:]:
      rest.text = ''
    paras[1].runs[0].text =subtitle
    for rest in paras[1].runs[1:]:
      rest.text = ''
    
    
    homeplayers = []
    awayplayers = []
    homeleaders = []
    awayleaders = []
    homeinfo = []
    awayinfo = []
    connection = pymysql.connect(host='localhost',
                                 user='root',
                                 password='961014',
                                 db=matchdb,
                                 charset='utf8mb4',
                                 cursorclass=pymysql.cursors.DictCursor)
    try:
      with connection.cursor() as cursor:
        sql = "SELECT * FROM Players WHERE Team = %s AND Valid = 1 ORDER BY KitNumber"
        cursor.execute(sql,hometeam)
        homeplayers = cursor.fetchall()
        cursor.execute(sql,awayteam)
        awayplayers = cursor.fetchall()
        sql = "SELECT * FROM Teams WHERE TeamName = %s"
        cursor.execute(sql,hometeam)
        homeinfo = cursor.fetchall()
        cursor.execute(sql,awayteam)
        awayinfo = cursor.fetchall()

    finally:
      connection.close()
    #print(len(tables))
    #for index,p in enumerate(tables):
    #  print('index:',index)
    #  for iq,q in enumerate(p.rows):
    #    print('row:',iq)
    #    for iw,w in enumerate(q.cells):
    #      print('cell',iw,":",w.text) 
    tables[0].rows[0].cells[0].paragraphs[0].runs[0].text = homeinfo[0]['TeamName']
    tables[0].rows[0].cells[7].paragraphs[0].runs[1].text += homeinfo[0]['KitColor']
    lenhc = len(tables[0].rows[0].cells[7].paragraphs[0].runs[0].text) + len(tables[0].rows[0].cells[7].paragraphs[0].runs[1].text)
    if lenhc >= 10:
      tables[0].rows[0].cells[7].paragraphs[0].runs[0].font.size = Pt(110/lenhc)
      tables[0].rows[0].cells[7].paragraphs[0].runs[1].font.size = Pt(110/lenhc)
      
      
    
    for i in range(len(homeplayers)):
      if not re.match('NANQI', matchdb) and not re.match('NYUQI', matchdb):
        tables[0].rows[i+2].cells[0].paragraphs[0].runs[0].text = str(homeplayers[i]['KitNumber'])
      if homeplayers[i]['ExtraInfo'] and not re.match('^\s+$',homeplayers[i]['ExtraInfo']):
        tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].text = homeplayers[i]['Name'] + '(' + homeplayers[i]['ExtraInfo'] + ')'
      else:
        tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].text = homeplayers[i]['Name']
      lenname = len(tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].text)
      utf8_lenname = len(tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].text.encode('utf-8'))
      lenname = (utf8_lenname - lenname) / 2 * 1.5 + lenname
      if lenname >= 16:
        tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].font.size = Pt(146/lenname)
      if homeplayers[i]['Suspension']:
        tables[0].rows[i+2].cells[1].paragraphs[0].runs[0].font.strike = True
        
     
    tables[0].rows[0].cells[10].paragraphs[0].runs[0].text = awayinfo[0]['TeamName']
    tables[0].rows[0].cells[16].paragraphs[0].runs[1].text += awayinfo[0]['KitColor']
    lenac = len(tables[0].rows[0].cells[16].paragraphs[0].runs[0].text) + len(tables[0].rows[0].cells[16].paragraphs[0].runs[1].text)
    if lenac >= 10:
      tables[0].rows[0].cells[16].paragraphs[0].runs[0].font.size = Pt(110/lenac)
      tables[0].rows[0].cells[16].paragraphs[0].runs[1].font.size = Pt(110/lenac)

    for i in range(len(awayplayers)):
      if not re.match('NANQI', matchdb):
        tables[0].rows[i+2].cells[10].paragraphs[0].runs[0].text = str(awayplayers[i]['KitNumber'])
      if awayplayers[i]['ExtraInfo'] and not re.match('\s+',awayplayers[i]['ExtraInfo']):
        tables[0].rows[i+2].cells[11].paragraphs[0].runs[0].text = awayplayers[i]['Name'] + '(' + awayplayers[i]['ExtraInfo'] + ')'
      else:
        tables[0].rows[i+2].cells[11].paragraphs[0].runs[0].text = awayplayers[i]['Name']
      lenname = len(tables[0].rows[i+2].cells[11].paragraphs[0].runs[0].text)
      utf8_lenname = len(tables[0].rows[i+2].cells[11].paragraphs[0].runs[0].text.encode('utf-8'))
      lenname = (utf8_lenname - lenname) / 2 * 1.5 + lenname
      if lenname >= 16:
        tables[0].rows[i+2].cells[11].paragraphs[0].runs[0].font.size = Pt(146/lenname)
      if awayplayers[i]['Suspension']:
        tables[0].rows[i+2].cells[11].paragraphs[0].runs[0].font.strike = True
      
    fname = '/var/www/TUFA/sheets/' + 'sheet.docx'
    doc.save(fname)
    return fname,fname[21:]


if __name__ == "__main__":
    create4subsheet('星期四晚上踢球吗', '清北人叉')
